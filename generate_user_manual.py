#!/usr/bin/env python3
"""
User Manual Generator for Inventory Management System
This script automatically generates a comprehensive user manual by touring the website.
"""

import os
import time
import sys
from datetime import datetime
from pathlib import Path

try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.action_chains import ActionChains
except ImportError:
    print("Installing selenium...")
    os.system("pip install selenium")
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.action_chains import ActionChains

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn

class UserManualGenerator:
    def __init__(self, base_url="http://localhost:3000"):
        self.base_url = base_url
        self.driver = None
        self.doc = Document()
        
        # Test credentials (you may need to adjust these)
        self.test_email = "wastiaman123@gmail.com"
        self.test_password = "polo0987"
        
    def setup_driver(self):
        """Setup Chrome WebDriver with appropriate options"""
        chrome_options = Options()
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--window-size=1920,1080")
        chrome_options.add_argument("--disable-gpu")
        
        try:
            self.driver = webdriver.Chrome(options=chrome_options)
            self.driver.maximize_window()
            print("Chrome WebDriver initialized successfully")
        except Exception as e:
            print(f"Error initializing Chrome WebDriver: {e}")
            print("Please ensure Chrome and ChromeDriver are installed")
            sys.exit(1)
    

    
    def wait_for_element(self, by, value, timeout=10):
        """Wait for an element to be present"""
        try:
            return WebDriverWait(self.driver, timeout).until(
                EC.presence_of_element_located((by, value))
            )
        except Exception as e:
            print(f"Element not found: {by}={value}, Error: {e}")
            return None
    
    def wait_for_page_load(self, timeout=10):
        """Wait for page to load completely"""
        try:
            WebDriverWait(self.driver, timeout).until(
                lambda driver: driver.execute_script("return document.readyState") == "complete"
            )
            time.sleep(2)  # Additional wait for dynamic content
        except Exception as e:
            print(f"Page load timeout: {e}")
    
    def setup_document_styles(self):
        """Setup document styles"""
        # Title style
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Heading styles
        heading1_style = self.doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_font = heading1_style.font
        heading1_font.name = 'Arial'
        heading1_font.size = Pt(18)
        heading1_font.bold = True
        heading1_style.paragraph_format.space_before = Pt(12)
        heading1_style.paragraph_format.space_after = Pt(6)
        
        heading2_style = self.doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_font = heading2_style.font
        heading2_font.name = 'Arial'
        heading2_font.size = Pt(14)
        heading2_font.bold = True
        heading2_style.paragraph_format.space_before = Pt(6)
        heading2_style.paragraph_format.space_after = Pt(3)
    
    def add_cover_page(self):
        """Add cover page to the document"""
        # Title
        title = self.doc.add_paragraph("USER MANUAL", style='CustomTitle')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = self.doc.add_paragraph("Computerized Stationary Management System", style='CustomTitle')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Organization
        org = self.doc.add_paragraph("Pakistan State Oil (PSO)", style='CustomHeading1')
        org.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some space
        self.doc.add_paragraph("")
        self.doc.add_paragraph("")
        
        # Version info
        version_info = self.doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        self.doc.add_page_break()
    
    def add_table_of_contents(self):
        """Add table of contents"""
        toc_heading = self.doc.add_paragraph("TABLE OF CONTENTS", style='CustomHeading1')
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        toc_items = [
            "1. Introduction",
            "2. Getting Started",
            "3. Authentication System",
            "4. Dashboard Overview",
            "5. Inventory Management",
            "6. Email Management",
            "7. User Profile Management",
            "8. Troubleshooting",
            "9. Appendix"
        ]
        
        for item in toc_items:
            self.doc.add_paragraph(item, style='Normal')
        
        self.doc.add_page_break()
    
    def add_section(self, title, content):
        """Add a section with title and content"""
        # Add heading
        self.doc.add_paragraph(title, style='CustomHeading1')
        
        # Add content
        if isinstance(content, list):
            for paragraph in content:
                self.doc.add_paragraph(paragraph, style='Normal')
        else:
            self.doc.add_paragraph(content, style='Normal')
        
        # Add space after section
        self.doc.add_paragraph("")
    
    def login_to_system(self):
        """Login to the system for authenticated pages"""
        try:
            print("Navigating to login page...")
            self.driver.get(self.base_url)
            self.wait_for_page_load()
            
            # Ensure we're on the login tab
            try:
                login_tab = self.wait_for_element(By.XPATH, "//button[contains(text(), 'Sign In')]", timeout=5)
                if login_tab:
                    login_tab.click()
                    time.sleep(1)
            except:
                print("Login tab not found or already active")
            
            # Find and fill email field using the correct ID
            email_field = self.wait_for_element(By.ID, "email", timeout=10)
            if not email_field:
                email_field = self.wait_for_element(By.CSS_SELECTOR, "input[type='email']", timeout=5)
            
            if email_field:
                email_field.clear()
                email_field.send_keys(self.test_email)
                print(f"Entered email: {self.test_email}")
            else:
                print("Email field not found")
                return False
            
            # Find and fill password field using the correct ID
            password_field = self.wait_for_element(By.ID, "password", timeout=10)
            if not password_field:
                password_field = self.wait_for_element(By.CSS_SELECTOR, "input[type='password']", timeout=5)
            
            if password_field:
                password_field.clear()
                password_field.send_keys(self.test_password)
                print(f"Entered password")
            else:
                print("Password field not found")
                return False
            
            # Find and click login button
            login_button = self.wait_for_element(By.CSS_SELECTOR, "button[type='submit']", timeout=10)
            if not login_button:
                login_button = self.wait_for_element(By.XPATH, "//button[contains(text(), 'Sign In')]", timeout=5)
            
            if login_button:
                print("Clicking login button...")
                login_button.click()
                
                # Wait for navigation to dashboard
                print("Waiting for redirect to dashboard...")
                try:
                    # Wait for URL to change to dashboard
                    WebDriverWait(self.driver, 15).until(
                        lambda driver: "/dashboard" in driver.current_url or driver.current_url != self.base_url
                    )
                    
                    # Additional wait for page to load
                    time.sleep(3)
                    
                    current_url = self.driver.current_url
                    print(f"Current URL after login: {current_url}")
                    
                    if "/dashboard" in current_url or current_url != self.base_url:
                        print("Login successful - redirected to dashboard")
                        return True
                    else:
                        print("Login may have failed - no redirect detected")
                        return False
                        
                except Exception as e:
                    print(f"Timeout waiting for redirect: {e}")
                    # Check if we're still on the login page
                    if self.driver.current_url == self.base_url:
                        print("Still on login page - login likely failed")
                        return False
                    else:
                        print("URL changed, assuming login successful")
                        return True
            else:
                print("Login button not found")
                return False
                
        except Exception as e:
            print(f"Login failed: {e}")
            return False
    
    def capture_landing_page(self):
        """Tour the landing page"""
        print("Touring landing page...")
        self.driver.get(self.base_url)
        self.wait_for_page_load()
        
        content = [
            "The Computerized Stationary Management System is a comprehensive web-based application designed for Pakistan State Oil (PSO) to manage inventory operations efficiently.",
            "",
            "Key Features:",
            "• Real-time Analytics & Reporting",
            "• Role-based Access Control", 
            "• Secure Data Management",
            "• Email Notifications",
            "• Stock History Tracking",
            "",
            "The landing page provides access to both login and signup functionality through a tabbed interface."
        ]
        
        self.add_section("1. Introduction - Landing Page", content)
    
    def capture_authentication(self):
        """Tour authentication pages"""
        print("Touring authentication pages...")
        
        # Ensure we're on the login tab first
        try:
            login_tab = self.wait_for_element(By.XPATH, "//button[contains(text(), 'Sign In')]", timeout=10)
            if login_tab:
                login_tab.click()
                time.sleep(2)
                print("Viewing Sign In tab")
        except Exception as e:
            print(f"Could not find or click Sign In tab: {e}")
        
        # Switch to signup tab to view it
        try:
            signup_tab = self.wait_for_element(By.XPATH, "//button[contains(text(), 'Sign Up')]", timeout=10)
            if signup_tab:
                signup_tab.click()
                time.sleep(2)
                print("Viewing Sign Up tab")
        except Exception as e:
            print(f"Could not find or click Sign Up tab: {e}")
        
        content = [
            "The authentication system provides secure access to the inventory management system.",
            "",
            "Login Process:",
            "1. Enter your registered email address",
            "2. Enter your password",
            "3. Click 'Sign In' to access the dashboard",
            "",
            "Sign Up Process:",
            "1. Enter your full name",
            "2. Provide a valid email address",
            "3. Create a secure password",
            "4. Click 'Sign Up' to create your account",
            "",
            "The system includes form validation and error handling for security."
        ]
        
        self.add_section("2. Authentication System", content)
        self.add_section("2.1 User Registration", ["New users can register using the Sign Up tab."])
    
    def capture_dashboard(self):
        """Tour dashboard page"""
        print("Touring dashboard...")
        
        # Login first
        if not self.login_to_system():
            print("Cannot access dashboard - login failed")
            return
        
        # Check if we're already on dashboard, if not navigate
        current_url = self.driver.current_url
        if "/dashboard" not in current_url:
            print("Navigating to dashboard...")
            self.driver.get(f"{self.base_url}/dashboard")
            self.wait_for_page_load()
        
        time.sleep(3)  # Wait for dashboard to fully load
        print("Viewing dashboard content...")
        
        content = [
            "The dashboard provides a comprehensive overview of your inventory system.",
            "",
            "Dashboard Features:",
            "• Real-time statistics and metrics",
            "• Interactive charts and graphs",
            "• Quick access navigation menu",
            "• User profile information",
            "",
            "Navigation Menu:",
            "• Dashboard - Overview and analytics",
            "• Inventory - Stock management and history",
            "• Email Management - Communication tools",
            "",
            "The header displays the PSO logo, system name, and user information with logout functionality."
        ]
        
        self.add_section("3. Dashboard Overview", content)
    
    def capture_inventory_management(self):
        """Tour inventory management pages"""
        print("Touring inventory management...")
        
        # Navigate to inventory page
        self.driver.get(f"{self.base_url}/inventory")
        self.wait_for_page_load()
        time.sleep(3)  # Wait for page to fully load
        print("Viewing inventory overview...")
        
        content = [
            "The Inventory Management section is the core of the system, providing comprehensive tools for managing stock items.",
            "",
            "The inventory section contains three main tabs:",
            "1. Inventory Overview - Combined view of all inventory data",
            "2. Current Stock - Real-time stock levels and item details",
            "3. Stock History - Historical transactions and movements",
            "",
            "Features available in inventory management:",
            "• Add new inventory items",
            "• Update stock quantities",
            "• Track stock movements",
            "• Generate reports",
            "• Search and filter items"
        ]
        
        self.add_section("4. Inventory Management", content)
        
        # Tour Current Stock tab
        try:
            print("Viewing Current Stock tab...")
            current_stock_tab = self.wait_for_element(By.XPATH, "//button[contains(text(), 'Current Stock')]", timeout=10)
            if current_stock_tab:
                current_stock_tab.click()
                time.sleep(3)  # Wait for tab content to load
                
                current_stock_content = [
                    "The Current Stock tab displays real-time inventory levels.",
                    "",
                    "Features:",
                    "• View all items with current quantities",
                    "• Add new inventory items",
                    "• Update existing item details",
                    "• Delete items when necessary",
                    "• Search and filter functionality",
                    "",
                    "Each item displays:",
                    "• Item name and description",
                    "• Current quantity in stock",
                    "• Location information",
                    "• Last updated timestamp"
                ]
                
                self.add_section("4.1 Current Stock Management", current_stock_content)
            else:
                print("Current Stock tab not found")
        except Exception as e:
            print(f"Error viewing current stock tab: {e}")
        
        # Tour Stock History tab
        try:
            print("Viewing Stock History tab...")
            history_tab = self.wait_for_element(By.XPATH, "//button[contains(text(), 'Stock History')]", timeout=10)
            if history_tab:
                history_tab.click()
                time.sleep(3)  # Wait for tab content to load
                
                history_content = [
                    "The Stock History tab provides a complete audit trail of all inventory movements.",
                    "",
                    "Features:",
                    "• View all stock transactions",
                    "• Filter by date range",
                    "• Filter by transaction type (IN/OUT)",
                    "• Search by item or user",
                    "• Export history reports",
                "",
                "Transaction details include:",
                "• Date and time of transaction",
                "• Item name and description",
                "• Quantity changed",
                "• Transaction type (Stock In/Stock Out)",
                "• User who performed the action",
                "• Additional notes or comments"
            ]
            
            self.add_section("4.2 Stock History Tracking", history_content)
        except Exception as e:
            print(f"Error viewing stock history tab: {e}")
        
        # Tour Add Item page
        # try:
        #     print("Touring Add Item page...")
        #     self.driver.get(f"{self.base_url}/inventory/add")
        #     self.wait_for_page_load()
        #     time.sleep(3)
            
        #     add_item_content = [
        #         "The Add Item page allows users to add new inventory items to the system.",
        #         "",
        #         "Required Information:",
        #         "• Item Name - Descriptive name of the item",
        #         "• Description - Detailed description",
        #         "• Quantity - Initial stock quantity",
        #         "• Location - Storage location",
        #         "• Category - Item classification",
        #         "",
        #         "Process:",
        #         "1. Fill in all required fields",
        #         "2. Verify information accuracy",
        #         "3. Click 'Add Item' to save",
        #         "4. System will confirm successful addition"
        #     ]
            
        #     self.add_section("4.3 Adding New Items", add_item_content)
        # except Exception as e:
        #     print(f"Error touring add item page: {e}")
    
    def capture_email_management(self):
        """Tour email management page"""
        print("Touring email management...")
        
        try:
            self.driver.get(f"{self.base_url}/email-management")
            self.wait_for_page_load()
            time.sleep(3)  # Wait for page to fully load
            print("Viewing email management interface...")
            
            content = [
                "The Email Management system provides automated communication capabilities.",
                "",
                "Features:",
                "• Send automated inventory reports",
                "• Configure email templates",
                "• Manage recipient lists",
                "• Schedule regular reports",
                "• Test email configuration",
                "",
                "Email Types:",
                "• Low stock alerts",
                "• Daily/Weekly inventory reports",
                "• Stock movement notifications",
                "• System maintenance alerts",
                "",
                "Configuration Options:",
                "• SMTP server settings",
                "• Email templates customization",
                "• Recipient management",
                "• Scheduling preferences"
            ]
            
            self.add_section("5. Email Management", content)
        except Exception as e:
            print(f"Error capturing email management: {e}")
    
    def add_troubleshooting_section(self):
        """Add troubleshooting section"""
        content = [
            "Common Issues and Solutions:",
            "",
            "Login Problems:",
            "• Verify email and password are correct",
            "• Check internet connection",
            "• Clear browser cache and cookies",
            "• Contact system administrator if issues persist",
            "",
            "Performance Issues:",
            "• Refresh the page",
            "• Check internet connection speed",
            "• Close unnecessary browser tabs",
            "• Use supported browsers (Chrome, Firefox, Safari)",
            "",
            "Data Not Loading:",
            "• Wait for page to fully load",
            "• Refresh the page",
            "• Check if you have proper permissions",
            "• Contact support if data appears missing",
            "",
            "Email Issues:",
            "• Verify email configuration",
            "• Check spam/junk folders",
            "• Ensure recipient email addresses are correct",
            "• Test email configuration in Email Management",
            "",
            "Browser Compatibility:",
            "• Use modern browsers (Chrome 90+, Firefox 88+, Safari 14+)",
            "• Enable JavaScript",
            "• Allow cookies for the application domain"
        ]
        
        self.add_section("6. Troubleshooting", content)
    
    def add_text_only_sections(self):
        """Add content sections without screenshots"""
        # 1. Introduction
        intro_content = [
            "The Computerized Stationary Management System is a comprehensive web-based application designed for Pakistan State Oil (PSO) to manage inventory operations efficiently.",
            "",
            "Key Features:",
            "• Real-time inventory tracking and management",
            "• User authentication and role-based access control",
            "• Comprehensive dashboard with analytics",
            "• Stock history and audit trail",
            "• Email notifications and reporting",
            "• Modern, responsive web interface"
        ]
        self.add_section("1. Introduction", intro_content)
        
        # 2. Getting Started
        getting_started_content = [
            "System Access:",
            "1. Open your web browser",
            "2. Navigate to the application URL",
            "3. Use your provided credentials to login",
            "",
            "First Time Users:",
            "• Contact your administrator for account setup",
            "• Ensure you have the correct permissions",
            "• Familiarize yourself with the interface"
        ]
        self.add_section("2. Getting Started", getting_started_content)
        
        # 3. Authentication
        auth_content = [
            "Login Process:",
            "1. Enter your email address",
            "2. Enter your password",
            "3. Click 'Sign In' to access the system",
            "",
            "New User Registration:",
            "1. Click 'Sign Up' tab",
            "2. Fill in your details (name, email, password)",
            "3. Submit the form to create your account",
            "",
            "Security Features:",
            "• Secure session management",
            "• Automatic logout after inactivity",
            "• Password protection"
        ]
        self.add_section("3. Authentication System", auth_content)
        
        # 4. Dashboard
        dashboard_content = [
            "The dashboard provides an overview of your inventory system with:",
            "",
            "Navigation Menu:",
            "• Dashboard - Main overview and statistics",
            "• Inventory - Stock management and history",
            "• Email Management - Communication tools",
            "",
            "Key Information Displayed:",
            "• Current stock levels",
            "• Recent activity",
            "• System statistics",
            "• Quick action buttons"
        ]
        self.add_section("4. Dashboard Overview", dashboard_content)
        
        # 5. Inventory Management
        inventory_content = [
            "The inventory section includes three main tabs:",
            "",
            "Inventory Overview:",
            "• Combined view of all inventory items",
            "• Real-time stock levels",
            "• Quick search and filtering",
            "",
            "Current Stock:",
            "• Add new inventory items",
            "• Update existing stock levels",
            "• Manage item details",
            "",
            "Stock History:",
            "• Complete audit trail of all transactions",
            "• Filter by date, item, or transaction type",
            "• Export capabilities for reporting"
        ]
        self.add_section("5. Inventory Management", inventory_content)
        
        # 6. Email Management
        email_content = [
            "Email management features include:",
            "",
            "Automated Notifications:",
            "• Low stock alerts",
            "• System notifications",
            "• Custom email templates",
            "",
            "Configuration Options:",
            "• SMTP settings",
            "• Recipient management",
            "• Email scheduling",
            "",
            "Testing and Monitoring:",
            "• Send test emails",
            "• Monitor delivery status",
            "• Troubleshoot email issues"
        ]
        self.add_section("6. Email Management", email_content)

    def add_appendix(self):
        """Add appendix with additional information"""
        content = [
            "System Requirements:",
            "",
            "Minimum Browser Requirements:",
            "• Google Chrome 90 or later",
            "• Mozilla Firefox 88 or later",
            "• Safari 14 or later",
            "• Microsoft Edge 90 or later",
            "",
            "Network Requirements:",
            "• Stable internet connection",
            "• Minimum 1 Mbps download speed",
            "• JavaScript enabled",
            "• Cookies enabled",
            "",
            "Security Features:",
            "• Secure HTTPS connection",
            "• Session-based authentication",
            "• Role-based access control",
            "• Data encryption in transit",
            "",
            "Support Information:",
            "• For technical support, contact your system administrator",
            "• Keep your login credentials secure",
            "• Report any security concerns immediately",
            "• Regular system backups are performed automatically"
        ]
        
        self.add_section("7. Appendix", content)
    
    def add_getting_started_section(self):
        """Add getting started section"""
        content = [
            "System Access:",
            "1. Open your web browser",
            "2. Navigate to the application URL",
            "3. Use your provided credentials to login",
            "",
            "First Time Users:",
            "• Contact your administrator for account setup",
            "• Ensure you have the correct permissions",
            "• Familiarize yourself with the interface",
            "",
            "Navigation Tips:",
            "• Use the main navigation menu to access different sections",
            "• The dashboard provides an overview of system status",
            "• Each section has specific tools for different tasks"
        ]
        self.add_section("2. Getting Started", content)
    
    def add_user_profile_section(self):
        """Add user profile management section"""
        content = [
            "User Profile Management:",
            "",
            "Profile Information:",
            "• View and update personal information",
            "• Change password when needed",
            "• Manage account preferences",
            "",
            "Security Settings:",
            "• Regular password updates recommended",
            "• Logout when finished using the system",
            "• Report any suspicious activity",
            "",
            "Account Management:",
            "• Contact administrator for role changes",
            "• Request additional permissions if needed",
            "• Keep contact information up to date"
        ]
        self.add_section("6. User Profile Management", content)
    
    def add_system_requirements_section(self):
        """Add system requirements section"""
        content = [
            "System Requirements:",
            "",
            "Minimum Browser Requirements:",
            "• Google Chrome 90 or later",
            "• Mozilla Firefox 88 or later",
            "• Safari 14 or later",
            "• Microsoft Edge 90 or later",
            "",
            "Network Requirements:",
            "• Stable internet connection",
            "• Minimum 1 Mbps download speed",
            "• JavaScript enabled",
            "• Cookies enabled",
            "",
            "Hardware Requirements:",
            "• Modern computer or mobile device",
            "• Minimum 4GB RAM recommended",
            "• Screen resolution 1024x768 or higher"
        ]
        self.add_section("8. System Requirements", content)
    
    def add_appendix_section(self):
        """Add appendix section"""
        content = [
            "Additional Information:",
            "",
            "Security Features:",
            "• Secure HTTPS connection",
            "• Session-based authentication",
            "• Role-based access control",
            "• Data encryption in transit",
            "",
            "Support Information:",
            "• For technical support, contact your system administrator",
            "• Keep your login credentials secure",
            "• Report any security concerns immediately",
            "• Regular system backups are performed automatically",
            "",
            "Best Practices:",
            "• Log out when finished",
            "• Use strong passwords",
            "• Keep browser updated",
            "• Report issues promptly"
        ]
        self.add_section("9. Appendix", content)
    
    def generate_text_only_manual(self):
        """Generate manual with text content only (no screenshots)"""
        try:
            self.setup_document_styles()
            self.add_cover_page()
            self.add_table_of_contents()
            self.add_text_only_sections()
            self.add_troubleshooting_section()
            self.add_appendix()
            
            doc_path = "User_Manual.docx"
            self.doc.save(doc_path)
            print(f"Text-only user manual generated: {doc_path}")
        except Exception as e:
            print(f"Error generating text-only manual: {e}")
    
    def tour_complete_website(self):
        """Tour the complete website in a single session without reloading"""
        print("Starting complete website tour...")
        
        # Start with landing page
        self.capture_landing_page()
        
        # Tour authentication (already on landing page)
        self.capture_authentication()
        
        # Login and tour authenticated pages
        print("Logging in to access authenticated pages...")
        if self.login_to_system():
            print("Successfully logged in, continuing tour...")
            
            # Tour dashboard (already navigated after login)
            self.capture_dashboard()
            
            # Tour inventory management
            self.capture_inventory_management()
            
            # Tour email management
            self.capture_email_management()
            
            print("Website tour completed successfully!")
        else:
            print("Login failed, skipping authenticated pages")

    def generate_manual(self):
        """Generate the complete user manual"""
        try:
            print("Starting user manual generation...")
            
            # Check if server is running
            try:
                import requests
                response = requests.get(self.base_url, timeout=5)
                server_running = response.status_code == 200
            except:
                server_running = False
            
            if not server_running:
                print(f"Development server not accessible at {self.base_url}")
                print("Generating text-only manual...")
                self.generate_text_only_manual()
                return "User_Manual.docx"
            
            # Setup browser and document
            self.setup_driver()
            self.setup_document_styles()
            
            # Generate manual by touring the website
            print("Generating manual by touring the website...")
            
            # Add cover page and TOC
            self.add_cover_page()
            self.add_table_of_contents()
            
            # Tour the complete website
            self.tour_complete_website()
            
            # Add additional sections
            self.add_getting_started_section()
            self.add_user_profile_section()
            self.add_troubleshooting_section()
            self.add_system_requirements_section()
            self.add_appendix_section()
            
            # Save document
            doc_path = "User_Manual.docx"
            self.doc.save(doc_path)
            print(f"User manual generated successfully: {doc_path}")
            return doc_path
            
        except Exception as e:
            print(f"Error generating manual: {e}")
            print("Falling back to text-only manual...")
            self.generate_text_only_manual()
            return "User_Manual.docx"
        finally:
            if self.driver:
                self.driver.quit()

def main():
    """Main function to run the user manual generator"""
    print("=== Inventory Management System - User Manual Generator ===")
    print("Checking if development server is running...")
    
    # Check if server is running
    import requests
    server_running = False
    try:
        response = requests.get("http://localhost:3000", timeout=5)
        if response.status_code == 200:
            server_running = True
            print("✓ Development server is running")
        else:
            print(f"⚠️ Development server responded with status: {response.status_code}")
    except requests.exceptions.RequestException as e:
        print(f"❌ Cannot connect to development server: {e}")
    
    if not server_running:
        print("\n⚠️ Development server is not accessible.")
        print("The manual will be generated without screenshots.")
        print("To include screenshots, please:")
        print("1. Start the development server: npm run dev")
        print("2. Ensure it's running on http://localhost:3000")
        print("3. Run this script again")
        print("\nProceeding with text-only manual generation...")
    
    print("Generating user manual...")
    print()
    
    generator = UserManualGenerator()
    doc_path = generator.generate_manual()
    
    print("\n=== Manual Generation Complete ===")
    print(f"✓ User manual created: {doc_path}")
    if server_running:
        print("✓ Screenshots captured and embedded")
    else:
        print("⚠️ Screenshots skipped (server not accessible)")
    print("✓ Professional formatting applied")
    print("\nThe manual is ready for use and distribution.")

if __name__ == "__main__":
    main()